import argparse
import json
import os
import sys
import zipfile
import subprocess

try:
    from vosk import Model, KaldiRecognizer
    from moviepy.editor import VideoFileClip
    from pydub import AudioSegment
    from deepmultilingualpunctuation import PunctuationModel
    from docx import Document
except ModuleNotFoundError:
    print("[INFO] Устанавливаю необходимые библиотеки...")
    subprocess.run([sys.executable, "-m", "pip", "install", "vosk", "moviepy", "pydub", "deepmultilingualpunctuation", "python-docx"], check=True)
    from vosk import Model, KaldiRecognizer
    from moviepy.editor import VideoFileClip
    from pydub import AudioSegment
    from deepmultilingualpunctuation import PunctuationModel
    from docx import Document


def install_ffmpeg():
    """Проверяет и устанавливает FFmpeg"""
    try:
        subprocess.run(["ffmpeg", "-version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
    except FileNotFoundError:
        print("[ERROR] FFmpeg не найден! Установите его вручную: https://ffmpeg.org/download.html")
        sys.exit(1)


def extract_audio(video_path, audio_path):
    """Извлекает аудио из видео и сохраняет в WAV"""
    video = VideoFileClip(video_path)
    video.audio.write_audiofile(audio_path, codec='pcm_s16le', fps=16000)


def transcribe_audio(audio_path, model_path):
    """Распознает речь и возвращает текст"""
    model = Model(model_path)
    recognizer = KaldiRecognizer(model, 16000)
    recognizer.SetWords(True)

    result_text = []
    audio = AudioSegment.from_wav(audio_path)
    audio = audio.set_channels(1).set_frame_rate(16000)

    step = 8000  # Блоки по 0.5 секунды
    for i in range(0, len(audio), step):
        segment = audio[i:i + step]
        if recognizer.AcceptWaveform(segment.raw_data):
            res = json.loads(recognizer.Result())
            result_text.append(res.get('text', ''))

    final_res = json.loads(recognizer.FinalResult())
    result_text.append(final_res.get('text', ''))

    return format_text_with_ai(' '.join(result_text))


def format_text_with_ai(text):
    """Форматирует текст с помощью нейросети и делает заглавные буквы в начале предложений"""
    model = PunctuationModel()
    formatted_text = model.restore_punctuation(text)
    
    # Разбиваем на предложения и делаем первую букву каждого предложения заглавной
    sentences = formatted_text.split('. ')
    formatted_sentences = [s.capitalize() for s in sentences]
    
    return '. '.join(formatted_sentences)


def save_to_docx(text, output_path):
    """Сохраняет текст в DOCX"""
    doc = Document()
    doc.add_heading("Розшифрований текст", level=1)

    paragraphs = text.split("\n\n")
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)

    doc.save(output_path)


def check_and_extract_model(model_path):
    """Проверяет и распаковывает модель Vosk, если она в zip"""
    if model_path.endswith(".zip"):
        extracted_folder = model_path.replace(".zip", "")
        if not os.path.exists(extracted_folder):
            print("[INFO] Распаковка модели Vosk...")
            with zipfile.ZipFile(model_path, 'r') as zip_ref:
                zip_ref.extractall(extracted_folder)
        return extracted_folder
    return model_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Транскрипция видео в текст (українською)")
    parser.add_argument("-video", required=True, help="Путь к видеофайлу")
    parser.add_argument("-out", required=True, help="Путь к выходному файлу .docx")
    parser.add_argument("-model", default="vosk-model-uk-v3", help="Путь к модели Vosk (или zip)")

    args = parser.parse_args()

    install_ffmpeg()
    model_path = check_and_extract_model(args.model)

    audio_temp = "temp_audio.wav"

    try:
        print("[INFO] Извлекаю аудио...")
        extract_audio(args.video, audio_temp)

        print("[INFO] Распознаю речь...")
        text = transcribe_audio(audio_temp, model_path)

        print("[INFO] Сохраняю в DOCX...")
        save_to_docx(text, args.out)

        print(f"[DONE] Файл сохранён: {args.out}")

    finally:
        if os.path.exists(audio_temp):
            os.remove(audio_temp)
